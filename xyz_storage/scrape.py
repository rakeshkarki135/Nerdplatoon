import re
import requests
from bs4 import BeautifulSoup
from docx import Document
import time

def scrap_blog(url, soup):
    try:
        # Find the <div> element with id "container"
        content = soup.find('div', class_="col-lg-8 col-md-7 blog-main-right-col")

        # If content is found, find the inner <div> with class "entry-content"
        if content:
            cont = content.find('div', class_="single-blog-section")
            return str(cont)  # Return the raw HTML of the div
        else:
            print(f"Content not found for {url}")
            return None
    except Exception as e:
        print(f"Error while scraping {url}: {e}")
        return None

# Main scraping function
def main():
    url = "https://qodesocial.com/tiktok-trends-riding-the-wave-of-viral-content-for-growth/"
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    print(soup)

    # Create a Document object to save the content
    doc = Document()
    doc.add_heading('Blog HTML Scraper (Raw <div> Content)', 0)

    # List of URLs to scrape
    urls = [
        "https://qodesocial.com/tiktok-trends-riding-the-wave-of-viral-content-for-growth/"
    ]

    # Loop through each URL
    for url in urls:
        print(f"Scraping HTML content from {url}")
        html_content = scrap_blog(url, soup)
        
        if html_content:
            # Add the HTML content to the Word document
            doc.add_heading(f'HTML Content from {url}', level=1)
            doc.add_paragraph(html_content)
            doc.add_paragraph("\n\n")  # Add spacing between different posts

        time.sleep(2)  # Sleep to avoid overloading the server with requests

    # Save the document to a .docx file
    doc.save('scraped_blog_html_content_div.docx')
    print("HTML content saved to scraped_blog_html_content_div.docx")

if __name__ == "__main__":
# Run the main function
    main()